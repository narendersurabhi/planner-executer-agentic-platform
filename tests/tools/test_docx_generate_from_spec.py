from __future__ import annotations

import json
from pathlib import Path

import pytest
from docx import Document

from libs.core.tool_registry import ToolExecutionError, ToolRegistry
from libs.tools.docx_generate import register_docx_tools


def _artifact_dir(tmp_path: Path) -> Path:
    root = tmp_path / "artifacts"
    root.mkdir(parents=True, exist_ok=True)
    return root


def _load_fixture(name: str) -> dict:
    path = Path("tests/fixtures") / name
    return json.loads(path.read_text(encoding="utf-8"))


def _make_registry() -> ToolRegistry:
    registry = ToolRegistry()
    register_docx_tools(registry)
    return registry


def test_generates_docx_file(tmp_path: Path, monkeypatch: pytest.MonkeyPatch) -> None:
    monkeypatch.setenv("ARTIFACTS_DIR", str(_artifact_dir(tmp_path)))
    registry = _make_registry()
    spec = _load_fixture("resume_ats_single_column_spec.json")
    payload = {"document_spec": spec, "path": "tests/resume.docx"}
    call = registry.execute("docx_generate_from_spec", payload, "id", "trace")
    assert call.status == "completed"
    output_path = Path(call.output_or_error["path"])
    assert output_path.exists()
    assert output_path.stat().st_size > 0


def test_path_traversal_blocked(tmp_path: Path, monkeypatch: pytest.MonkeyPatch) -> None:
    monkeypatch.setenv("ARTIFACTS_DIR", str(_artifact_dir(tmp_path)))
    registry = _make_registry()
    spec = _load_fixture("resume_ats_single_column_spec.json")
    with pytest.raises(ToolExecutionError):
        registry.get("docx_generate_from_spec").handler(
            {"document_spec": spec, "path": "../evil.docx"}
        )


def test_invalid_validation_report_blocks_render(
    tmp_path: Path, monkeypatch: pytest.MonkeyPatch
) -> None:
    monkeypatch.setenv("ARTIFACTS_DIR", str(_artifact_dir(tmp_path)))
    registry = _make_registry()
    spec = _load_fixture("resume_ats_single_column_spec.json")
    with pytest.raises(ToolExecutionError, match="document_spec validation failed"):
        registry.get("docx_generate_from_spec").handler(
            {
                "document_spec": spec,
                "path": "tests/blocked.docx",
                "validation_report": {
                    "valid": False,
                    "errors": [
                        {
                            "path": "/blocks/0/text",
                            "message": "text/paragraph requires text: string",
                        }
                    ],
                },
            }
        )


def test_missing_path_auto_derives_output_path() -> None:
    import tempfile

    from pytest import MonkeyPatch

    monkeypatch = MonkeyPatch()
    monkeypatch.setenv("ARTIFACTS_DIR", str(_artifact_dir(Path(tempfile.mkdtemp()))))
    registry = _make_registry()
    spec = _load_fixture("resume_ats_single_column_spec.json")
    try:
        call = registry.execute("docx_generate_from_spec", {"document_spec": spec}, "id", "trace")
        assert call.status == "completed"
        assert call.output_or_error["path"].endswith(".docx")
    finally:
        monkeypatch.undo()


def test_strict_unresolved_placeholder_raises(
    tmp_path: Path, monkeypatch: pytest.MonkeyPatch
) -> None:
    monkeypatch.setenv("ARTIFACTS_DIR", str(_artifact_dir(tmp_path)))
    registry = _make_registry()
    spec = {
        "tokens": {"name": "A"},
        "blocks": [{"type": "paragraph", "text": "{{missing}}"}],
    }
    with pytest.raises(ToolExecutionError):
        registry.get("docx_generate_from_spec").handler(
            {"document_spec": spec, "path": "tests/strict.docx", "strict": True}
        )


def test_repeat_expansion(tmp_path: Path, monkeypatch: pytest.MonkeyPatch) -> None:
    monkeypatch.setenv("ARTIFACTS_DIR", str(_artifact_dir(tmp_path)))
    registry = _make_registry()
    spec = {
        "tokens": {
            "experience": [
                {"company": "Alpha"},
                {"company": "Beta"},
            ]
        },
        "blocks": [
            {
                "type": "repeat",
                "items": "{{experience}}",
                "as": "r",
                "template": [
                    {"type": "paragraph", "text": "{{r.company}}"},
                ],
            }
        ],
    }
    call = registry.execute(
        "docx_generate_from_spec",
        {"document_spec": spec, "path": "tests/repeat.docx"},
        "id",
        "trace",
    )
    assert call.status == "completed"
    doc = Document(call.output_or_error["path"])
    text = "\n".join(paragraph.text for paragraph in doc.paragraphs)
    assert "Alpha" in text
    assert "Beta" in text


def test_inline_emphasis_markers_render_as_bold_and_italic(
    tmp_path: Path, monkeypatch: pytest.MonkeyPatch
) -> None:
    monkeypatch.setenv("ARTIFACTS_DIR", str(_artifact_dir(tmp_path)))
    registry = _make_registry()
    spec = {
        "blocks": [
            {
                "type": "paragraph",
                "text": "Perfect. Below is a <b>high quality Codex prompt</b> with _inline emphasis_.",
            }
        ]
    }
    call = registry.execute(
        "docx_generate_from_spec",
        {"document_spec": spec, "path": "tests/inline_emphasis.docx"},
        "id",
        "trace",
    )
    assert call.status == "completed"
    doc = Document(call.output_or_error["path"])
    runs = doc.paragraphs[0].runs
    assert any(run.text == "high quality Codex prompt" and run.bold for run in runs)
    assert any(run.text == "inline emphasis" and run.italic for run in runs)


def test_cover_letter_style_spacing(tmp_path: Path, monkeypatch: pytest.MonkeyPatch) -> None:
    monkeypatch.setenv("ARTIFACTS_DIR", str(_artifact_dir(tmp_path)))
    registry = _make_registry()
    spec = {
        "blocks": [
            {"type": "paragraph", "style": "cover_letter_name", "text": "Narender Surabhi"},
            {"type": "paragraph", "style": "cover_letter_date", "text": "February 13, 2026"},
            {"type": "paragraph", "style": "cover_letter_recipient", "text": "Hiring Team"},
            {"type": "paragraph", "style": "cover_letter_salutation", "text": "Dear Hiring Team,"},
            {"type": "paragraph", "style": "cover_letter_body", "text": "Paragraph."},
            {"type": "paragraph", "style": "cover_letter_closing", "text": "Sincerely,"},
            {"type": "paragraph", "style": "cover_letter_signature", "text": "Narender Surabhi"},
        ],
    }
    call = registry.execute(
        "docx_generate_from_spec",
        {"document_spec": spec, "path": "tests/coverletter_spacing.docx"},
        "id",
        "trace",
    )
    assert call.status == "completed"
    doc = Document(call.output_or_error["path"])
    paragraphs = doc.paragraphs
    assert len(paragraphs) == 7

    name_run = paragraphs[0].runs[0]
    assert name_run.bold is True
    assert paragraphs[0].paragraph_format.space_after.pt == pytest.approx(3.0)

    assert paragraphs[1].paragraph_format.space_before.pt == pytest.approx(10.0)
    assert paragraphs[1].paragraph_format.space_after.pt == pytest.approx(8.0)
    assert paragraphs[2].paragraph_format.space_after.pt == pytest.approx(2.0)
    assert paragraphs[3].paragraph_format.space_before.pt == pytest.approx(8.0)
    assert paragraphs[3].paragraph_format.space_after.pt == pytest.approx(8.0)
    assert paragraphs[4].paragraph_format.space_after.pt == pytest.approx(8.0)
    assert paragraphs[5].paragraph_format.space_before.pt == pytest.approx(8.0)
    assert paragraphs[5].paragraph_format.space_after.pt == pytest.approx(2.0)
    assert paragraphs[6].paragraph_format.space_after.pt == pytest.approx(0.0)


def test_experience_heading_and_group_spacing(
    tmp_path: Path, monkeypatch: pytest.MonkeyPatch
) -> None:
    monkeypatch.setenv("ARTIFACTS_DIR", str(_artifact_dir(tmp_path)))
    registry = _make_registry()
    spec = {
        "blocks": [
            {"type": "heading", "level": 1, "style": "section_heading", "text": "EXPERIENCE"},
            {"type": "paragraph", "style": "role_title", "text": "Software Engineer, AI/ML"},
            {
                "type": "paragraph",
                "style": "role_meta",
                "text": "Acentra Health | Dec 2016 - Present",
            },
            {
                "type": "paragraph",
                "style": "role_group_heading",
                "text": "Applied AI & LLM Systems",
            },
            {"type": "bullets", "items": ["Built platform X"]},
        ],
    }
    call = registry.execute(
        "docx_generate_from_spec",
        {"document_spec": spec, "path": "tests/experience_spacing.docx"},
        "id",
        "trace",
    )
    assert call.status == "completed"
    doc = Document(call.output_or_error["path"])
    paragraphs = doc.paragraphs
    assert len(paragraphs) >= 5

    assert paragraphs[0].paragraph_format.keep_with_next is True
    assert paragraphs[1].paragraph_format.space_before.pt == pytest.approx(10.0)
    assert paragraphs[1].paragraph_format.keep_with_next is True
    assert paragraphs[2].paragraph_format.space_after.pt == pytest.approx(4.0)
    assert paragraphs[2].paragraph_format.keep_with_next is True
    assert paragraphs[3].paragraph_format.space_before.pt == pytest.approx(6.0)
    assert paragraphs[3].paragraph_format.keep_with_next is True
