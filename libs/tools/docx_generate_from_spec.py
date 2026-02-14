from __future__ import annotations

import os
import re
from pathlib import Path
from typing import Any, Dict, Iterable, List

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph

from libs.core.models import RiskLevel, ToolIntent, ToolSpec


def _artifacts_dir() -> Path:
    return Path(os.getenv("ARTIFACTS_DIR", "/shared/artifacts"))


def register_docx_tools(registry) -> None:
    from libs.core.tool_registry import Tool

    registry.register(
        Tool(
            spec=ToolSpec(
                name="docx_generate_from_spec",
                description="Generate an ATS-friendly .docx document from a DocumentSpec JSON",
                usage_guidance=(
                    "Provide path (relative .docx filename). document_spec is "
                    "resolved from memory (document_spec:latest) unless explicitly provided. "
                    "path can be resolved from memory (docx_path:latest) when using derive_output_filename. "
                    "render_context (merged into tokens), and strict (default true). "
                    "Supported blocks: text, paragraph, heading, bullets, spacer, "
                    "optional_paragraph, repeat. "
                    "If you only have an output folder (e.g., output_dir in job.context_json), "
                    "derive a safe filename and pass the full relative path as path."
                ),
                input_schema={
                    "type": "object",
                    "properties": {
                        "document_spec": {"type": "object"},
                        "path": {"type": "string"},
                        "render_context": {"type": "object"},
                        "strict": {"type": "boolean"},
                    },
                },
                output_schema={
                    "type": "object",
                    "properties": {
                        "path": {"type": "string"},
                        "bytes_written": {"type": "integer"},
                    },
                    "required": ["path", "bytes_written"],
                },
                memory_reads=["job_context", "task_outputs"],
                timeout_s=25,
                risk_level=RiskLevel.medium,
                tool_intent=ToolIntent.render,
            ),
            handler=_docx_generate_from_spec,
        )
    )


def _docx_generate_from_spec(payload: Dict[str, Any]) -> Dict[str, Any]:
    document_spec = payload.get("document_spec")
    if not isinstance(document_spec, dict):
        _tool_error(
            "document_spec missing (not found in memory). Provide document_spec explicitly."
        )

    path = payload.get("path")
    if not isinstance(path, str) or not path:
        _tool_error("path must be a non-empty string")
    if not path.endswith(".docx"):
        _tool_error("path must end with .docx")
    if Path(path).is_absolute():
        _tool_error("path must be relative to /shared/artifacts")

    render_context = payload.get("render_context", {})
    if render_context is None:
        render_context = {}
    if not isinstance(render_context, dict):
        _tool_error("render_context must be an object")

    strict = payload.get("strict", True)
    if not isinstance(strict, bool):
        _tool_error("strict must be a boolean")

    output_path = _safe_join_artifacts_path(path)
    output_path.parent.mkdir(parents=True, exist_ok=True)

    theme = document_spec.get("theme", {})
    tokens = document_spec.get("tokens", {})
    if tokens is None:
        tokens = {}
    if not isinstance(tokens, dict):
        _tool_error("document_spec.tokens must be an object")

    merged_tokens = {**tokens, **render_context}
    context = {"tokens": merged_tokens, **merged_tokens}

    document = Document()
    _apply_theme(document, theme if isinstance(theme, dict) else {})

    blocks = document_spec.get("blocks", [])
    if not isinstance(blocks, list):
        _tool_error("document_spec.blocks must be an array")

    _render_blocks(document, blocks, context, strict)

    document.save(output_path)
    bytes_written = output_path.stat().st_size
    return {"path": str(output_path), "bytes_written": int(bytes_written)}


def _safe_join_artifacts_path(path: str) -> Path:
    artifacts_dir = _artifacts_dir()
    artifacts_dir.mkdir(parents=True, exist_ok=True)
    candidate = (artifacts_dir / path).resolve()
    root = artifacts_dir.resolve()
    if not str(candidate).startswith(str(root) + os.sep) and candidate != root:
        _tool_error("Invalid path outside /shared/artifacts")
    return candidate


def _apply_theme(document: Document, theme: Dict[str, Any]) -> None:
    fonts = theme.get("fonts", {}) if isinstance(theme.get("fonts", {}), dict) else {}
    font_sizes = (
        theme.get("font_sizes", {}) if isinstance(theme.get("font_sizes", {}), dict) else {}
    )
    body_font = fonts.get("body", "Calibri")
    body_size = font_sizes.get("body", 11)

    normal_style = document.styles["Normal"]
    normal_style.font.name = body_font
    normal_style.font.size = Pt(float(body_size))
    spacing = theme.get("spacing", {}) if isinstance(theme.get("spacing", {}), dict) else {}
    para_after = spacing.get("para_after_pt")
    if para_after is not None:
        normal_style.paragraph_format.space_after = Pt(float(para_after))
    line_spacing = spacing.get("line")
    if line_spacing is not None:
        normal_style.paragraph_format.line_spacing = float(line_spacing)

    heading_size = font_sizes.get("h1")
    heading_before = spacing.get("heading_before_pt", 12)
    tight_after = spacing.get("tight_after_heading_pt", 4)
    try:
        heading_style = document.styles["Heading 1"]
        if heading_size is not None:
            heading_style.font.size = Pt(float(heading_size))
        heading_style.font.bold = True
        if heading_before is not None:
            heading_style.paragraph_format.space_before = Pt(float(heading_before))
        if tight_after is not None:
            heading_style.paragraph_format.space_after = Pt(float(tight_after))
    except KeyError:
        pass

    margins = (
        theme.get("page_margins_in", {})
        if isinstance(theme.get("page_margins_in", {}), dict)
        else {}
    )
    section = document.sections[0]
    for key, attr in (
        ("top", "top_margin"),
        ("bottom", "bottom_margin"),
        ("left", "left_margin"),
        ("right", "right_margin"),
    ):
        value = margins.get(key)
        if value is not None:
            section.__setattr__(attr, Inches(float(value)))


def _next_rendered_block(
    blocks: List[Dict[str, Any]], start_index: int, context: Dict[str, Any], strict: bool
) -> Dict[str, Any] | None:
    for idx in range(start_index + 1, len(blocks)):
        candidate = blocks[idx]
        if not isinstance(candidate, dict):
            continue
        block_type = candidate.get("type")
        if block_type == "spacer":
            continue
        if block_type == "optional_paragraph":
            when_value = candidate.get("when")
            if not _evaluate_condition(when_value, context, strict):
                continue
        return candidate
    return None


def _render_blocks(
    document: Document, blocks: Iterable[Dict[str, Any]], context: Dict[str, Any], strict: bool
) -> None:
    block_list = [block for block in blocks if isinstance(block, dict)]
    for index, block in enumerate(block_list):
        if not isinstance(block, dict):
            continue

        block_type = block.get("type")

        if block_type in {"text", "paragraph"}:
            text = _render_text(str(block.get("text", "")), context, strict)
            paragraph = _add_paragraph(document, text, block.get("style"))
            if block.get("style") == "contact" and paragraph is not None:
                next_block = _next_rendered_block(block_list, index, context, strict)
                next_is_contact = (
                    isinstance(next_block, dict)
                    and next_block.get("style") == "contact"
                    and next_block.get("type") in {"text", "paragraph", "optional_paragraph"}
                )
                paragraph.paragraph_format.space_before = Pt(0)
                paragraph.paragraph_format.space_after = Pt(0 if next_is_contact else 10)

        elif block_type == "heading":
            level = block.get("level", 1)
            if not isinstance(level, int) or level < 1 or level > 3:
                level = 1
            text = _render_text(str(block.get("text", "")), context, strict)
            paragraph = document.add_heading(text, level=level)
            if block.get("style") == "section_heading":
                paragraph.paragraph_format.space_before = Pt(12)
                paragraph.paragraph_format.space_after = Pt(4)
                paragraph.paragraph_format.keep_with_next = True
                _set_paragraph_bottom_border(paragraph)

        elif block_type == "bullets":
            items = _resolve_items(block.get("items"), context, strict)
            for item in items:
                text = _render_text(str(item), context, strict)
                _add_paragraph(document, text, block.get("style"), bullet=True)

        elif block_type == "spacer":
            document.add_paragraph("")

        elif block_type == "optional_paragraph":
            when_value = block.get("when")
            should_render = _evaluate_condition(when_value, context, strict)
            if should_render:
                text = _render_text(str(block.get("text", "")), context, strict)
                paragraph = _add_paragraph(document, text, block.get("style"))
                if block.get("style") == "contact" and paragraph is not None:
                    next_block = _next_rendered_block(block_list, index, context, strict)
                    next_is_contact = (
                        isinstance(next_block, dict)
                        and next_block.get("style") == "contact"
                        and next_block.get("type") in {"text", "paragraph", "optional_paragraph"}
                    )
                    paragraph.paragraph_format.space_before = Pt(0)
                    paragraph.paragraph_format.space_after = Pt(0 if next_is_contact else 10)

        elif block_type == "repeat":
            items = _resolve_items(block.get("items"), context, strict)
            alias = block.get("as", "item")
            template = block.get("template", [])
            if not isinstance(alias, str) or not alias:
                alias = "item"
            if not isinstance(template, list):
                continue
            for item in items:
                scoped_context = dict(context)
                scoped_context[str(alias)] = item
                _render_blocks(document, template, scoped_context, strict)

        else:
            # Ignore unknown blocks for robustness
            continue


def _add_paragraph(
    document: Document, text: str, style_hint: Any, bullet: bool = False
) -> Paragraph | None:
    paragraph = document.add_paragraph()
    if bullet:
        paragraph.style = "List Bullet"
    if style_hint == "section_rule":
        paragraph.paragraph_format.space_before = Pt(2)
        paragraph.paragraph_format.space_after = Pt(6)
        _set_paragraph_bottom_border(paragraph)
        return paragraph
    if style_hint == "term_def":
        paragraph.paragraph_format.space_after = Pt(2)
        _add_term_definition_runs(paragraph, text)
        return paragraph
    if style_hint == "role_title":
        run = paragraph.add_run(text)
        run.bold = True
        paragraph.paragraph_format.space_before = Pt(10)
        paragraph.paragraph_format.space_after = Pt(1)
        paragraph.paragraph_format.keep_with_next = True
        return paragraph
    if style_hint == "role_meta":
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(4)
        paragraph.paragraph_format.keep_with_next = True
        _add_role_meta_runs(paragraph, text)
        return paragraph
    if style_hint == "role_group_heading":
        run = paragraph.add_run(text)
        run.bold = True
        run.italic = True
        paragraph.paragraph_format.space_before = Pt(6)
        paragraph.paragraph_format.space_after = Pt(2)
        paragraph.paragraph_format.keep_with_next = True
        paragraph.paragraph_format.keep_together = True
        return paragraph
    if style_hint == "role_header":
        _set_right_tab_stop(paragraph, document)
        if "\t" in text:
            left, right = text.split("\t", 1)
            left_run = paragraph.add_run(left.strip())
            left_run.bold = True
            paragraph.add_run("\t")
            paragraph.add_run(right.strip())
        else:
            run = paragraph.add_run(text)
            run.bold = True
        paragraph.paragraph_format.space_before = Pt(4)
        paragraph.paragraph_format.space_after = Pt(2)
        return paragraph

    run = paragraph.add_run(text)
    if style_hint == "cover_letter_name":
        run.bold = True
        run.font.size = Pt(20)
        run.font.name = "Calibri"
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(3)
        return paragraph
    if style_hint == "cover_letter_date":
        paragraph.paragraph_format.space_before = Pt(10)
        paragraph.paragraph_format.space_after = Pt(8)
        return paragraph
    if style_hint == "cover_letter_recipient":
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(2)
        return paragraph
    if style_hint == "cover_letter_salutation":
        paragraph.paragraph_format.space_before = Pt(8)
        paragraph.paragraph_format.space_after = Pt(8)
        return paragraph
    if style_hint == "cover_letter_body":
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(8)
        return paragraph
    if style_hint == "cover_letter_closing":
        paragraph.paragraph_format.space_before = Pt(8)
        paragraph.paragraph_format.space_after = Pt(2)
        return paragraph
    if style_hint == "cover_letter_signature":
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)
        return paragraph
    if style_hint == "title":
        run.bold = True
        run.font.size = Pt(22)
        run.font.name = "Calibri"
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(3)
    if style_hint == "subtitle":
        run.bold = False
        run.font.size = Pt(13)
        run.font.name = "Calibri"
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(6)
    if style_hint in {"body_bold"}:
        run.bold = True
    if style_hint == "contact":
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.paragraph_format.space_before = Pt(0)
        run.font.name = "Calibri"
        run.font.size = Pt(10.5)
    if style_hint == "dates_right":
        paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    if style_hint == "divider":
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.space_before = Pt(2)
        paragraph.paragraph_format.space_after = Pt(2)
        run.font.size = Pt(8)
    if bullet:
        paragraph.paragraph_format.space_after = Pt(2)
        paragraph.paragraph_format.left_indent = Pt(18)
        paragraph.paragraph_format.first_line_indent = Pt(-9)
    return paragraph


def _set_right_tab_stop(paragraph, document: Document) -> None:
    try:
        section = document.sections[0]
        usable_width = section.page_width - section.left_margin - section.right_margin
        paragraph.paragraph_format.tab_stops.add_tab_stop(
            usable_width, alignment=WD_TAB_ALIGNMENT.RIGHT
        )
    except Exception:
        return


def _set_paragraph_bottom_border(paragraph) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "999999")
    p_bdr.append(bottom)


def _add_term_definition_runs(paragraph, text: str) -> None:
    if ":" in text:
        term, definition = text.split(":", 1)
        term_run = paragraph.add_run(term.strip() + ":")
        term_run.bold = True
        if definition.strip():
            paragraph.add_run(" " + definition.strip())
        return
    if " - " in text:
        term, definition = text.split(" - ", 1)
        term_run = paragraph.add_run(term.strip())
        term_run.bold = True
        if definition.strip():
            paragraph.add_run(" - " + definition.strip())
        return
    run = paragraph.add_run(text)
    run.bold = True


def _add_role_meta_runs(paragraph, text: str) -> None:
    if "|" in text:
        left, right = text.split("|", 1)
        paragraph.add_run(left.strip())
        paragraph.add_run(" | ")
        right_run = paragraph.add_run(right.strip())
        right_run.italic = True
        return
    paragraph.add_run(text)


def _resolve_items(value: Any, context: Dict[str, Any], strict: bool) -> List[Any]:
    if isinstance(value, list):
        return value
    if isinstance(value, str):
        resolved = _resolve_placeholder(value, context, strict, return_raw=True)
        if isinstance(resolved, list):
            return resolved
        if strict:
            _tool_error("items must resolve to an array")
        return [resolved]
    if value is None:
        return []
    if strict:
        _tool_error("items must be an array")
    return [value]


def _evaluate_condition(value: Any, context: Dict[str, Any], strict: bool) -> bool:
    if isinstance(value, bool):
        return value
    if isinstance(value, str):
        resolved = _resolve_placeholder(value, context, strict, return_raw=True)
        return bool(resolved)
    if value is None:
        return False
    if strict:
        _tool_error("optional_paragraph.when must be a boolean or placeholder")
    return bool(value)


def _render_text(text: str, context: Dict[str, Any], strict: bool) -> str:
    def replace(match: re.Match[str]) -> str:
        expr = match.group(1).strip()
        if not expr:
            if strict:
                _tool_error("Empty placeholder")
            return match.group(0)
        value = _resolve_path(expr, context, strict)
        if value is None:
            if strict:
                _tool_error(f"Unresolved placeholder: {expr}")
            return match.group(0)
        # If lists show up in-line, join them
        if isinstance(value, list):
            return ", ".join(str(x) for x in value)
        return str(value)

    rendered = re.sub(r"\{\{\s*(.*?)\s*\}\}", replace, text)
    # Normalize en-dash/em-dash to hyphen for ATS-friendly output.
    rendered = rendered.replace("\u2013", "-").replace("\u2014", "-")
    if strict and re.search(r"\{\{\s*(.*?)\s*\}\}", rendered):
        _tool_error("Unresolved placeholders remain after rendering")
    return rendered


def _resolve_placeholder(text: str, context: Dict[str, Any], strict: bool, return_raw: bool) -> Any:
    matches = re.findall(r"\{\{\s*(.*?)\s*\}\}", text)
    if len(matches) == 1 and text.strip() == f"{{{{{matches[0]}}}}}":
        return _resolve_path(matches[0].strip(), context, strict)
    rendered = _render_text(text, context, strict)
    return rendered if return_raw else rendered


def _resolve_path(expr: str, context: Dict[str, Any], strict: bool) -> Any:
    tokens = _tokenize_path(expr)
    current: Any = context
    for token in tokens:
        if isinstance(token, int):
            if isinstance(current, list) and 0 <= token < len(current):
                current = current[token]
                continue
            if strict:
                _tool_error(f"Index out of range: {token} for {expr}")
            return None
        # token is a string key
        if isinstance(current, dict) and token in current:
            current = current[token]
            continue
        if strict:
            _tool_error(f"Unknown placeholder: {expr}")
        return None
    return current


def _tokenize_path(expr: str) -> List[Any]:
    parts: List[Any] = []
    for match in re.finditer(r"([^[.\]]+)|(\[(\d+)\])", expr):
        key = match.group(1)
        index = match.group(3)
        if key is not None:
            parts.append(key)
        elif index is not None:
            parts.append(int(index))
    return parts


def _tool_error(message: str) -> None:
    from libs.core.tool_registry import ToolExecutionError

    raise ToolExecutionError(message)
